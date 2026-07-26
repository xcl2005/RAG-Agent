"""文档加载模块：把不同格式文件解析成纯文本。

RAG 的第一步是 ingestion，也就是“资料入库”。
不同文件格式的解析方法不同，所以单独放在 loaders.py：
- txt/md：直接读取文本。
- pdf：逐页提取文本，保留 page 元数据，方便引用来源。
- docx：提取段落。
- html：去掉 script/style/nav/footer 后提取正文。
"""

from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from rag_agent.schemas import RawDocument

# 项目支持的文件后缀。上传接口也会复用这个集合做校验。
SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}
MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def iter_files(path: str | Path) -> list[Path]:
    """返回需要导入的文件列表。

    如果 path 是单个文件，直接返回它；
    如果 path 是目录，递归扫描目录下所有支持格式的文件。
    """

    root = Path(path)
    if not root.exists():
        raise FileNotFoundError(f"ingest path does not exist: {root}")
    if root.is_file():
        if root.suffix.lower() not in SUPPORTED_SUFFIXES:
            raise ValueError(f"unsupported file type: {root.suffix.lower()}")
        return [root]
    return sorted(p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES)


def load_text_file(path: Path) -> list[RawDocument]:
    """加载 txt / markdown。

    errors="ignore" 可以跳过少量编码异常，避免一个坏字符导致整个导入失败。
    """

    text = path.read_text(encoding="utf-8-sig", errors="replace")
    base_metadata = {
        "source": str(path),
        "title": path.name,
        "type": path.suffix.lower(),
    }
    if path.suffix.lower() != ".md":
        return [RawDocument(text=text, metadata=base_metadata)]

    # Heading-aware Markdown units preserve section names in citations and keep
    # unrelated chapters from being blended before the recursive chunker runs.
    docs: list[RawDocument] = []
    heading = path.stem
    buffer: list[str] = []
    for line in text.splitlines(keepends=True):
        match = MARKDOWN_HEADING_RE.match(line.rstrip("\r\n"))
        if match and buffer:
            docs.append(
                RawDocument(
                    text="".join(buffer),
                    metadata={**base_metadata, "heading": heading},
                )
            )
            buffer = []
        if match:
            heading = match.group(2).strip()
        buffer.append(line)
    if buffer:
        docs.append(
            RawDocument(
                text="".join(buffer),
                metadata={**base_metadata, "heading": heading},
            )
        )
    return [doc for doc in docs if doc.text.strip()]


def load_pdf(path: Path) -> list[RawDocument]:
    """加载 PDF。

    这里按页切成 RawDocument，而不是整个 PDF 一起处理。
    好处是最终引用可以显示 page，面试时可以说“支持来源溯源”。
    """

    reader = PdfReader(str(path))
    docs: list[RawDocument] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                RawDocument(
                    text=text,
                    metadata={
                        "source": str(path),
                        "title": path.name,
                        "type": ".pdf",
                        "page": i,
                    },
                )
            )
    return docs


def load_docx(path: Path) -> list[RawDocument]:
    """加载 Word 文档，只提取普通段落文本。

    注意：复杂表格、图片 OCR 没有做。这是简化版项目，避免难度过高。
    """

    doc = DocxDocument(str(path))
    blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                blocks.append(" | ".join(values))
    text = "\n".join(blocks)
    return [
        RawDocument(
            text=text,
            metadata={"source": str(path), "title": path.name, "type": ".docx"},
        )
    ]


def load_html(path: Path) -> list[RawDocument]:
    """加载 HTML，并尽量只保留正文。"""

    html_text = path.read_text(encoding="utf-8-sig", errors="replace")
    soup = BeautifulSoup(html_text, "lxml")

    # 网页里的脚本、样式、导航、页脚通常不是知识内容，会污染检索。
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()

    text = soup.get_text("\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title = soup.title.get_text(strip=True) if soup.title else path.name
    return [
        RawDocument(
            text="\n".join(lines),
            metadata={"source": str(path), "title": title, "type": path.suffix.lower()},
        )
    ]


def load_document(path: str | Path) -> list[RawDocument]:
    """根据后缀分发到对应解析函数。"""

    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in {".txt", ".md"}:
        return load_text_file(p)
    if suffix == ".pdf":
        return load_pdf(p)
    if suffix == ".docx":
        return load_docx(p)
    if suffix in {".html", ".htm"}:
        return load_html(p)
    raise ValueError(f"Unsupported file type: {p}")


def load_documents(path: str | Path) -> list[RawDocument]:
    """加载一个文件或目录，返回所有 RawDocument。"""

    docs: list[RawDocument] = []
    for file in iter_files(path):
        docs.extend(load_document(file))
    return docs
