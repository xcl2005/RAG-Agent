from types import SimpleNamespace

import pytest
from docx import Document

from rag_agent.ingest import loaders


def test_iter_files_filters_supported_formats_and_rejects_bad_inputs(tmp_path):
    (tmp_path / "guide.md").write_text("# Guide", encoding="utf-8")
    (tmp_path / "notes.txt").write_text("notes", encoding="utf-8")
    (tmp_path / "image.png").write_bytes(b"png")

    assert [path.name for path in loaders.iter_files(tmp_path)] == ["guide.md", "notes.txt"]
    with pytest.raises(FileNotFoundError):
        loaders.iter_files(tmp_path / "missing")
    with pytest.raises(ValueError, match="unsupported"):
        loaders.iter_files(tmp_path / "image.png")


def test_markdown_loader_preserves_headings(tmp_path):
    path = tmp_path / "handbook.md"
    path.write_text(
        "# Access Control\nUse least privilege.\n\n## Rotation\nRotate keys regularly.\n",
        encoding="utf-8",
    )

    docs = loaders.load_document(path)

    assert [doc.metadata["heading"] for doc in docs] == ["Access Control", "Rotation"]
    assert docs[0].metadata["title"] == "handbook.md"
    assert "least privilege" in docs[0].text


def test_html_loader_removes_navigation_and_scripts(tmp_path):
    path = tmp_path / "page.html"
    path.write_text(
        """
        <html><head><title>Runbook</title><style>.x{}</style></head>
        <body><nav>menu</nav><main>Restart the worker.</main>
        <script>alert('x')</script><footer>copyright</footer></body></html>
        """,
        encoding="utf-8",
    )

    doc = loaders.load_document(path)[0]

    assert doc.metadata["title"] == "Runbook"
    assert "Restart the worker." in doc.text
    assert "alert" not in doc.text
    assert "menu" not in doc.text


def test_docx_loader_includes_paragraphs_and_tables(tmp_path):
    path = tmp_path / "policy.docx"
    document = Document()
    document.add_paragraph("Policy summary")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Platform"
    document.save(path)

    doc = loaders.load_document(path)[0]

    assert "Policy summary" in doc.text
    assert "Owner | Platform" in doc.text
    assert doc.metadata["type"] == ".docx"


def test_pdf_loader_retains_only_non_empty_pages(monkeypatch, tmp_path):
    class Page:
        def __init__(self, text):
            self.text = text

        def extract_text(self):
            return self.text

    monkeypatch.setattr(
        loaders,
        "PdfReader",
        lambda _: SimpleNamespace(pages=[Page("first page"), Page(None), Page("third page")]),
    )
    path = tmp_path / "paper.pdf"
    path.write_bytes(b"%PDF-test")

    docs = loaders.load_pdf(path)

    assert [doc.metadata["page"] for doc in docs] == [1, 3]
    assert [doc.text for doc in docs] == ["first page", "third page"]


def test_load_documents_and_unsupported_dispatch(tmp_path):
    (tmp_path / "a.txt").write_text("A", encoding="utf-8")
    (tmp_path / "b.htm").write_text("<main>B</main>", encoding="utf-8")

    docs = loaders.load_documents(tmp_path)

    assert len(docs) == 2
    with pytest.raises(ValueError, match="Unsupported"):
        loaders.load_document(tmp_path / "data.csv")
